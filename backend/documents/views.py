import os
import json
import csv
import tempfile
import docx
import pdfplumber
import requests
import time
import mimetypes
import pickle
from io import BytesIO, StringIO
from bs4 import BeautifulSoup
from PIL import Image
from gtts import gTTS
from django.conf import settings
from rest_framework.views import APIView 
from rest_framework.response import Response
from rest_framework import status, permissions

# ---- Google OAuth Drive ----
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload,MediaIoBaseDownload,MediaInMemoryUpload
from google.auth.transport.requests import Request

# ---- Gemini ----
from google import genai
from google.genai import types
from google.api_core import exceptions

# ---- Models & Serializers ----
from .models import Document, SummarizationSession, SummarizationMessage
from .serializers import DocumentSerializer, SummarizationSessionSerializer, SummarizationMessageSerializer



def upload_file_to_drive(file_path_or_buffer, filename, mimetype='application/octet-stream'):
    """
    Uploads a file (from a local path or an in-memory buffer) to Google Drive.
    """
    service = get_drive_service()
    parent_folder_id = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
    if not parent_folder_id:
        raise Exception("GOOGLE_DRIVE_FOLDER_ID is not set in .env")

    metadata = {"name": filename, "parents": [parent_folder_id]}
    
    if isinstance(file_path_or_buffer, str): # It's a file path
        media = MediaFileUpload(file_path_or_buffer, mimetype=mimetype, resumable=True)
    else: # It's an in-memory buffer (BytesIO)
        media = MediaInMemoryUpload(file_path_or_buffer.read(), mimetype=mimetype)

    file = service.files().create(
        body=metadata, media_body=media, fields="id, webViewLink, webContentLink"
    ).execute()
    return file
# ===========================================================
# 🧠 TEXT EXTRACTION
# ===========================================================
def extract_text_from_file(file_content_stream, file_name):
    """
    Extracts text directly from an in-memory binary stream (BytesIO)
    for all supported file types.
    """
    ext = os.path.splitext(file_name)[-1].lower()
    text = ""
    try:
        if ext == ".pdf":
            # pdfplumber can read directly from a file-like object
            with pdfplumber.open(file_content_stream) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        elif ext == ".docx":
            # python-docx can read directly from a file-like object
            doc = docx.Document(file_content_stream)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        # For text-based formats, first decode the binary stream into a string
        elif ext in [".csv", ".json", ".html", ".htm", ".xml", ".txt"]:
            decoded_content = file_content_stream.read().decode('utf-8', errors='ignore')

            if ext == ".csv":
                # Use StringIO to treat the string as a file for the csv reader
                csv_file = StringIO(decoded_content)
                reader = csv.reader(csv_file)
                for row in reader:
                    text += ", ".join(row) + "\n"
            elif ext == ".json":
                # Parse the JSON string
                data = json.loads(decoded_content)
                text = json.dumps(data, indent=2)
            elif ext in [".html", ".htm", ".xml"]:
                # Parse the HTML/XML string
                soup = BeautifulSoup(decoded_content, "html.parser")
                text = soup.get_text(separator="\n")
            else: # .txt
                text = decoded_content
        
        elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"]:
            text = " OCR not implemented yet for images."

        else:
            return f" Unsupported file type: {ext}"

        char_count = len(text.strip())
        print(f"--- DEBUG: Extracted {char_count} characters from {file_name} ---")
        if char_count == 0:
            return " Could not extract any readable text."

        return text.strip()
    except Exception as e:
        return f" ERROR processing file in memory: {e}"
# ===========================================================
# 📤 GOOGLE DRIVE UPLOAD (OAuth)
# ===========================================================
def get_drive_service():
    """
    Authenticate using OAuth (client_secret.json + token.pickle)
    Returns Drive service object.
    """
    creds = None
    client_secret = os.getenv("GOOGLE_DRIVE_CLIENT_SECRET_FILE")
    token_path = os.path.join(settings.BASE_DIR, "token.pickle")
    scopes = ["https://www.googleapis.com/auth/drive.file"]

    # Load saved credentials if available
    if os.path.exists(token_path):
        with open(token_path, "rb") as token:
            creds = pickle.load(token)

    # Refresh or login if needed
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(client_secret, scopes)
            creds = flow.run_local_server(port=0)
            with open(token_path, "wb") as token:
                pickle.dump(creds, token)

    return build("drive", "v3", credentials=creds)

def extract_text(file_content_stream, file_name):
    """
    Extracts text directly from an in-memory binary stream (BytesIO)
    for all supported file types.
    """
    ext = os.path.splitext(file_name)[-1].lower()
    text = ""
    try:
        if ext == ".pdf":
            # pdfplumber can read directly from a file-like object
            with pdfplumber.open(file_content_stream) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        elif ext == ".docx":
            # python-docx can read directly from a file-like object
            doc = docx.Document(file_content_stream)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        # For text-based formats, first decode the binary stream into a string
        elif ext in [".csv", ".json", ".html", ".htm", ".xml", ".txt"]:
            decoded_content = file_content_stream.read().decode('utf-8', errors='ignore')

            if ext == ".csv":
                # Use StringIO to treat the string as a file for the csv reader
                csv_file = StringIO(decoded_content)
                reader = csv.reader(csv_file)
                for row in reader:
                    text += ", ".join(row) + "\n"
            elif ext == ".json":
                # Parse the JSON string
                data = json.loads(decoded_content)
                text = json.dumps(data, indent=2)
            elif ext in [".html", ".htm", ".xml"]:
                # Parse the HTML/XML string
                soup = BeautifulSoup(decoded_content, "html.parser")
                text = soup.get_text(separator="\n")
            else: # .txt
                text = decoded_content
        
        elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"]:
            text = " OCR not implemented yet for images."

        else:
            return f" Unsupported file type: {ext}"

        char_count = len(text.strip())
        print(f"--- DEBUG: Extracted {char_count} characters from {file_name} ---")
        if char_count == 0:
            return " Could not extract any readable text."

        return text.strip()
    except Exception as e:
        return f" ERROR processing file in memory: {e}"


# --- API VIEWS ---

class DocumentUploadView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, *args, **kwargs):
        serializer = DocumentSerializer(data=request.data, context={"request": request})
        if not serializer.is_valid():
            return Response(serializer.errors, status=400)

        # ✅ Save file locally first (Django default)
        document = serializer.save(user=request.user)

        try:
            # --- 1. Prepare local path and metadata ---
            local_path = getattr(document.file, "path", None)
            if not local_path or not os.path.exists(local_path):
                return Response({"error": "Local file not found after upload."}, status=400)

            filename = f"user_{request.user.id}_{os.path.basename(local_path)}"
            folder_id = os.getenv("GOOGLE_DRIVE_FOLDER_ID")

            # --- 2. Initialize Drive Service ---
            service = get_drive_service()
            if not service:
                raise RuntimeError("Google Drive service not initialized. Check credentials.")

            # --- 3. Upload to Google Drive ---
            media = MediaFileUpload(local_path, resumable=True)
            metadata = {"name": filename}
            if folder_id:
                metadata["parents"] = [folder_id]

            drive_file = (
                service.files()
                .create(
                    body=metadata,
                    media_body=media,
                    fields="id, webViewLink, webContentLink",
                )
                .execute()
            )

            # --- 4. Save Drive metadata in DB ---
            document.drive_file_id = drive_file.get("id")
            document.file_url = drive_file.get("webViewLink")
            document.web_content_link = drive_file.get("webContentLink")
            document.save(update_fields=["drive_file_id", "file_url", "web_content_link"])

            # --- 5. Delete local file after successful upload ---
            try:
                os.remove(local_path)
                print(f"✅ DRIVE INFO: Uploaded '{filename}' and removed local copy.")
            except Exception as del_err:
                print(f"⚠️ Warning: Uploaded but could not delete local file: {del_err}")

        except Exception as e:
            print(f"❌ DRIVE UPLOAD FAILED: {e}")
            # File stays locally if upload fails (safe fallback)

        # ✅ Return document data (even if Drive upload failed)
        return Response(DocumentSerializer(document).data, status=201)

    permission_classes = [permissions.IsAuthenticated]
    def post(self, request, *args, **kwargs):
        serializer = DocumentSerializer(data=request.data, context={"request": request})
        if not serializer.is_valid(): return Response(serializer.errors, status=400)
        document = serializer.save(user=request.user)
        try:
            local_path = document.file.path
            filename = f"user_{request.user.id}_{os.path.basename(local_path)}"
            
            service = get_drive_service()
            folder_id = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
            media = MediaFileUpload(local_path, resumable=True)
            metadata = {"name": filename, "parents": [folder_id]}
            drive_file = service.files().create(body=metadata, media_body=media, fields="id, webViewLink, webContentLink").execute()
            
            document.drive_file_id = drive_file.get("id")
            document.file_url = drive_file.get("webViewLink")
            document.web_content_link = drive_file.get("webContentLink")
            document.save()

            os.remove(local_path)
            print(f"DRIVE INFO: Successfully uploaded '{filename}' and removed local copy.")

        except Exception as e:
            print(f"DRIVE UPLOAD FAILED: {e}")
            pass
        return Response(DocumentSerializer(document).data, status=201)


class SummarizeView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        file_ids = request.data.get("files", [])
        if not isinstance(file_ids, list):
            return Response({"error": "files must be a list of IDs"}, status=400)

        docs = Document.objects.filter(id__in=file_ids, user=request.user)
        if not docs.exists():
            return Response({"error": "No documents found for this user."}, status=404)

        drive_service = get_drive_service()
        combined_text = ""

        for doc in docs:
            text = ""
            try:
                if doc.drive_file_id and drive_service:
                    print(f"Summarizer: Processing Google Drive file: {doc.file.name}")
                    gdrive_request = drive_service.files().get_media(fileId=doc.drive_file_id)
                    file_content_stream = BytesIO()
                    downloader = MediaIoBaseDownload(file_content_stream, gdrive_request)
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()
                    file_content_stream.seek(0)
                    text = extract_text(file_content_stream, doc.file.name)
                else:
                    # fallback: local file
                    print(f"Summarizer: Processing local file: {doc.file.path}")
                    with open(doc.file.path, 'rb') as f:
                        file_content_stream = BytesIO(f.read())
                    text = extract_text(file_content_stream, doc.file.name)
            except Exception as e:
                text = f"⚠️ ERROR extracting text from {doc.file.name}: {e}"
            combined_text += text + "\n\n"

        if not combined_text.strip():
            return Response({"error": "No readable text could be extracted from the document(s)."}, status=400)

        # --- Gemini Summarization ---
        try:
            client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
            prompt = f"""
You are a professional document analyst. 
Summarize the following document(s) into a **highly detailed, well-structured Markdown report**.

⚠️ IMPORTANT: 
- Use clear markdown headers: `### 1. Overview`, `### 2. Important Details`, etc. 
- Always use bullet points (`- ...`) for lists, never long paragraphs. 
- Highlight key terms/dates/names in **bold**.
- Add line breaks between sections for readability.

Your output MUST strictly follow this structure:

### 1. Overview
(2–4 sentences max, in plain text.)

### 2. Important Details
- **Clause/Instruction** → explanation
- **Date/Name/Number** → explanation
- (Continue listing EVERYTHING important)

### 3. Context & Purpose
- Why the document exists
- Who it is for
- How it is used

### 4. Implications
- **Rule broken** → consequence
- **Missed requirement** → penalty

### 5. Extra Observations
- Errors, missing parts, inconsistencies
- Anything unusual or noteworthy

### 6. Verbatim Quotes
- "Copy key phrases here"
- "Use exact wording from the text"

---
📄 Document Content:
{combined_text[:12000]}
"""

            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[{"role": "user", "parts": [{"text": prompt}]}]
            )
            summary_text = getattr(response, "text", None)
            if not summary_text:
                return Response({"error": "Gemini returned no summary."}, status=500)

            session = SummarizationSession.objects.create(
                user=request.user,
                document=docs.first(),
                title=f'Summary of "{os.path.basename(docs.first().file.name)}"',
                summary_text=summary_text,
            )

            return Response({
                "session_id": session.id,
                "title": session.title,
                "summary": summary_text,
                "created_at": session.created_at.isoformat(),
            }, status=200)

        except exceptions.GoogleAPICallError as e:
            return Response({"error": f"Gemini API Error: {e.message}"}, status=500)
        except Exception as e:
            return Response({"error": f"Unexpected error: {str(e)}"}, status=500)
# ===========================================================
# 🧾 LIST SUMMARIES
# ===========================================================
class SummarizeListView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        sessions = SummarizationSession.objects.filter(user=request.user).order_by("-created_at")
        return Response(SummarizationSessionSerializer(sessions, many=True).data, status=200)

# ===========================================================
# 💬 CHAT WITH SUMMARY
# ===========================================================
class SummarizeChatView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, session_id):
        query = request.data.get("query", "").strip()
        if not query:
            return Response({"error": "Empty query"}, status=400)

        try:
            session = SummarizationSession.objects.get(id=session_id, user=request.user)
        except SummarizationSession.DoesNotExist:
            return Response({"error": "Session not found"}, status=404)

        prompt = f"""
Context:
{session.summary_text}

User Question:
{query}
"""
        client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])]
        )
        answer = getattr(response, "text", None) or "⚠️ No response."

        SummarizationMessage.objects.create(session=session, role="user", content=query)
        SummarizationMessage.objects.create(session=session, role="assistant", content=answer)
        return Response({"reply": answer}, status=200)

# ===========================================================
# 🔊 AUDIO SUMMARY
# ===========================================================
class AudioSummarizeView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, session_id):
        lang = request.data.get("language", "en")
        try:
            session = SummarizationSession.objects.get(id=session_id, user=request.user)
            
            # 1. Prepare narration text (clean up markdown)
            narration = session.summary_text.replace("**", "").replace("#", "").strip()
            if not narration:
                return Response({"error": "Summary text is empty, cannot generate audio."}, status=400)

            # 2. Generate the TTS audio directly into an in-memory buffer
            print(f"--- AUDIO: Generating TTS for session {session_id} in '{lang}' ---")
            audio_buffer = BytesIO()
            tts = gTTS(narration, lang=lang)
            tts.write_to_fp(audio_buffer)
            audio_buffer.seek(0)
            print("--- AUDIO: TTS generated in memory. ---")

            # 3. Upload the in-memory audio file to Google Drive
            drive_filename = f"audio_summary_user_{request.user.id}_session_{session_id}_{lang}.mp3"
            print(f"--- AUDIO: Uploading '{drive_filename}' to Google Drive... ---")
            
            drive_file = upload_file_to_drive(audio_buffer, drive_filename, mimetype='audio/mpeg')
            drive_link = drive_file.get("webViewLink") # Use the view link for browser playback
            
            print(f"--- AUDIO: Upload successful. Link: {drive_link} ---")

            # 4. Return the public Google Drive URL
            return Response({
                "audio_url": drive_link,
                "narration": narration # Still useful to send back for the UI
            }, status=200)

        except SummarizationSession.DoesNotExist:
            return Response({"error": "Session not found"}, status=404)
        except Exception as e:
            return Response({"error": f"Failed to generate audio summary: {str(e)}"}, status=500)

